"""Testes do adapter real DoclingChunker (multi-formato).

Carrega o Docling + tokenizer (mais lento). Marcado como 'integration'.
O teste de PDF usa o PDF de exemplo do repositório (data/pdf/...).
"""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from rag_service.adapters.outbound.docling_chunker import (
    DoclingChunker,
    UnsupportedFormatError,
)
from rag_service.domain.models import RawDocument

pytestmark = pytest.mark.integration

HTML = """
<html><body>
  <h1>Café</h1>
  <p>A origem do café na Etiópia é contada pela lenda do pastor Kaldi.</p>
  <h2>Tipos</h2>
  <table>
    <tr><th>Espécie</th><th>Sabor</th></tr>
    <tr><td>Arábica</td><td>suave</td></tr>
    <tr><td>Robusta</td><td>forte</td></tr>
  </table>
</body></html>
"""

PDF_PATH = (
    Path(__file__).resolve().parents[2]
    / "data"
    / "pdf"
    / "cafe-uma-breve-historia-de-sua-origem-consumo-cultura-diferenciacoes-e-utilidades-gastronomicas20.pdf"
)


@pytest.fixture(scope="module")
def chunker():
    return DoclingChunker()


def html_doc(html: str) -> RawDocument:
    return RawDocument(filename="document.html", data=html.encode("utf-8"))


def test_html_produces_chunks_with_text(chunker):
    chunks = chunker.chunk(html_doc(HTML))

    assert chunks
    assert all(c.text for c in chunks)
    assert all(c.chunk_id.startswith("chunk-") for c in chunks)


def test_html_table_is_markdown(chunker):
    texto = " ".join(c.text for c in chunker.chunk(html_doc(HTML)))
    assert "Arábica" in texto
    assert "|" in texto


def test_preserves_utf8_accents(chunker):
    html = "<html><body><p>café, ótimo à noite, coração e ação</p></body></html>"
    text = " ".join(c.text for c in chunker.chunk(html_doc(html)))

    assert "café" in text
    assert "coração" in text
    assert "Ã©" not in text and "Ã§" not in text


def test_empty_html_returns_no_chunks(chunker):
    assert chunker.chunk(html_doc("<html><body></body></html>")) == []


def test_markdown_is_extracted(chunker):
    md = "# Café\n\nA origem do café na Etiópia é contada pela lenda de Kaldi.\n"
    raw = RawDocument(filename="notas.md", data=md.encode("utf-8"))

    chunks = chunker.chunk(raw)

    assert chunks
    todos = " ".join(c.contextualized_text.lower() for c in chunks)
    assert "café" in todos and "kaldi" in todos


def test_docx_is_extracted(chunker):
    doc = DocxDocument()
    doc.add_heading("Café", level=1)
    doc.add_paragraph("A origem do café na Etiópia é contada pela lenda de Kaldi.")
    buffer = BytesIO()
    doc.save(buffer)

    raw = RawDocument(filename="relatorio.docx", data=buffer.getvalue())
    chunks = chunker.chunk(raw)

    assert chunks
    assert "Kaldi" in " ".join(c.text for c in chunks)


def test_unsupported_format_is_rejected(chunker):
    raw = RawDocument(filename="planilha.xlsx", data=b"conteudo")

    with pytest.raises(UnsupportedFormatError):
        chunker.chunk(raw)


@pytest.mark.skipif(not PDF_PATH.exists(), reason="PDF de exemplo ausente")
def test_pdf_is_extracted_and_chunked(chunker):
    raw = RawDocument(filename=PDF_PATH.name, data=PDF_PATH.read_bytes())

    chunks = chunker.chunk(raw)

    assert len(chunks) > 5  # o PDF tem várias páginas
    assert all(c.text for c in chunks)
    todos = " ".join(c.contextualized_text.lower() for c in chunks)
    assert "café" in todos
